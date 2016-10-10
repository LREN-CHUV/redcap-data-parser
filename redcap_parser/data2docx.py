#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  3 17:24:32 2016

@author: guillaume
"""

from docx import Document
import redcap_parser as rp
from tree import Tree
from docx.shared import Pt


def one_line_2_docx(p,name,value=''):
    """
        Prints to word document format:

            <b>name</b>value\n

    """
    # print component
    p.add_run(name).bold = True
    if value != '':
        p.add_run(value)
        p.add_run('\n')



def comp2docx(document,comp,bSorted=False):


    p = document.add_paragraph()

    p.add_run('----------------------------------------------------------------------------------------------------------------------------------------------------\n')
    one_line_2_docx(p,'Component(' + comp.subject_ID  + ') \n')
    one_line_2_docx(p,'Task leader: ', comp.task_leader      )

    if bSorted == False:
        for task in comp.tasks:
            print 'processing task: ', task.name
            task2docx(document,task)
    else:

        comp.sort()
        leaf_nodes = comp.tree.get_leaf_nodes()


        for leaf in leaf_nodes:

            # PRINT HEADING OF ONE LEAF

            elements = leaf.path
            i = 1
            for elem in elements:
                document.add_heading(elem, level=i)
                i = i + 1

            p.add_run('\n')

            for idx in leaf.idx:

                task2docx(document,comp.tasks[idx],b_print_heading=False)

def sorted_task2docx(document,tasks):

    tree = Tree()
    for i in range(0,len(tasks)):
        if len(tasks[i].build_block_belong) == 3:
            tree.add_element(tasks[i].build_block_belong,i)


    p = document.add_paragraph()


    leaf_nodes = tree.get_leaf_nodes()
    for leaf in leaf_nodes:
                # PRINT HEADING OF ONE LEAF

            elements = leaf.path
            i = 1
            for elem in elements:
                document.add_heading(elem, level=i)
                i = i + 1

            p.add_run('\n')

            for idx in leaf.idx:
                task2docx(document,tasks[idx],b_print_heading=False)

    return tree

def task2docx(document,task,b_print_heading=True):
    
    bAddRelease = False
    bAddUser    = False

    # To which building block your component belongs to

    if b_print_heading:
        elements = task.build_block_belong # ('SOFTWARE', 'Data Factory (DF)', 'Feature Engineering')
        i = 1;
        for elem in elements:
            document.add_heading(elem, level=i)
            i = i + 1

    bPrint = True

    if bPrint == True:
        #Prior Dependencies table:
        p = document.add_paragraph()

        task_name   = task.name
        task_number = task.task_number
        task_desc   = task.short_desc_comp

        comp_dep_desc2docx(document,task_name,task_number,task_desc)

        p = document.add_paragraph()

        one_line_2_docx(p,'Dependencies:')

        #  add table


        data      = task.build_block_data
        soft      = task.build_block_soft
        serv      = task.build_block_serv
        model     = task.build_block_model
        report    = task.build_block_reports

        dependencies2docx(document,data,soft,serv,model,report)

        
        if bAddRelease:
            # Prior Release table:
            p = document.add_paragraph()
            p.add_run('\n')
            one_line_2_docx(p,'Releases:')
            releases2docx(document,task.planned_functionality)

        if bAddUser:
            # Prior User table:
            p = document.add_paragraph()
            p.add_run('\n')
            one_line_2_docx(p,'User and Use cases:')
            usecase2docx(document,task.users,task.short_desc_usecase)

def  comp_dep_desc2docx(wdoc,task_name,task_number,short_desc_comp):
    table       = wdoc.add_table(rows=3, cols=2)
    table.style = 'Normal Table'
    
    if isinstance(short_desc_comp,unicode):
        short_desc_comp = str(short_desc_comp.encode('ascii','ignore'))
        
    if isinstance(task_number,unicode):
        task_number = str(task_number.encode('ascii','ignore')) 
        
    if isinstance(task_number,unicode):
        task_name = str(task_name.encode('ascii','ignore'))         
        
        
    if not isinstance(short_desc_comp,basestring):
        print short_desc_comp
        short_desc_comp = 'invalid'
        
        
    if short_desc_comp is None:
        short_desc_comp = 'None'
        
    
    short_desc_comp = (short_desc_comp.decode("utf-8")).encode('ascii','ignore')         

        
    print task_name, ' type: ', type(short_desc_comp), ' short_desc: ',  short_desc_comp

    row = table.rows[0]
    row.cells[0].paragraphs[0].add_run('Component').bold = True
    row.cells[1].text = task_name

    row = table.rows[1]
    row.cells[0].paragraphs[0].add_run('Description').bold = True
    row.cells[1].text = task_number

    row = table.rows[2]
    row.cells[0].paragraphs[0].add_run('Contribution task').bold = True
    row.cells[1].text = short_desc_comp


def add_dependency2table(table,name,data,row_idx):
    
    if len(data.keys()) == 0:
        return row_idx
    
    if row_idx < len(table.rows):
       
        # Software Data
        row = table.rows[row_idx]
        row.cells[0].paragraphs[0].add_run(name).bold = True

        idx = row_idx
        for key in data:
            row = table.rows[idx]
            soft_desc = ', '.join(data[key])
            #row.cells[1].text = key
            row.cells[1].text = key + ': ' + soft_desc
            idx = idx + 1
        row_idx = idx
        
    else:
        print '[Warning data2docx.py] r_idx >= len(table.rows):   r_idx: ', row_idx, '  len(table.rows):  ', len(table.rows)
        
    return row_idx     


def dependencies2docx(wdoc,data,soft,services,models,report):
    """

    Takes Data and software information and puts it into a table

    Prameters
    ---------

    wdoc: python-docx word document object
    data_hosp: DATA > Hospital Data > [] list
    data_ref : DATA > Reference Data > [] list
    soft     : dict
    services : dict

    """
    # get number of data 
    num_data        = len(data.keys())

    # get number of sofware
    num_soft        = len(soft.keys())

    # get number of services
    num_serv        = len(services.keys())
    
    # get number of models
    num_models      = len(models.keys())

    tab_num_rows = 0
    tab_num_rows = tab_num_rows + num_data + num_soft + num_serv + num_models
    
    print ' ========== '
    print 'num_data:   ', num_data
    print 'num_soft:   ', num_soft    
    print 'num_serv:   ', num_serv    
    print 'num_models: ', num_models
    

    if tab_num_rows == 0:
        return

    table       = wdoc.add_table(rows=tab_num_rows, cols=2)
    table.style = 'TableGrid'
    
    r_idx = 0    
    
    r_idx = add_dependency2table(table,'DATA',data,r_idx)
    r_idx = add_dependency2table(table,'SOFTWARE',soft,r_idx)
    r_idx = add_dependency2table(table,'SERVICES',services,r_idx)
    r_idx = add_dependency2table(table,'MODELS',models,r_idx)
    r_idx = add_dependency2table(table,'REPORTS',report,r_idx)


def releases2docx(wdoc,release):
    """
        Takes release information  (Planned functional etc..)


    """

    num_releases = len(release.keys())

    table = wdoc.add_table(rows=num_releases, cols=2)
    table.style = 'TableGrid'



    data = []
    for key in release:
        
        name = key[-2:]
        data.append([name,release[key]])


    data = sorted(data)

    idx = 0
    for d in data:
        row = table.rows[idx]
        row.cells[0].text = 'Planned functionality M' + d[0]
        row.cells[1].text = d[1]
        idx = idx + 1
         

    #
    #row.cells[1].text = release[key]


def usecase2docx(wdoc,users,short_desc_usecase):


    table = wdoc.add_table(rows=2, cols=2)
    table.style = 'TableGrid'


    # Heading
    row = table.rows[0]
    row.cells[0].text = 'User'
    row.cells[1].text = 'Use cases'

    # User & Use cases
    row = table.rows[1]
    row.cells[0].text = users
    row.cells[1].text = short_desc_usecase


def summary_comp2docx(summary_data,stype=0):
    """
    Parameters
    ----------
    
    wdoc: Document object
    summary_data: list of lists.


    """    


    num_rows     = len(summary_data)

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(6)
    
    document.add_paragraph()
    
    
    table       = document.add_table(rows=num_rows+1, cols=4)
    table.style = 'TableGrid'

    row = table.rows[0]
    if stype == 0:
        row.cells[0].paragraphs[0].add_run('Build').bold = True
        row.cells[1].paragraphs[0].add_run('Product').bold = True
        row.cells[2].paragraphs[0].add_run('Component').bold = True
        row.cells[3].paragraphs[0].add_run('Task').bold = True

    elif stype == 1:

        row.cells[0].paragraphs[0].add_run('Component').bold = True
        row.cells[1].paragraphs[0].add_run('Planned functionalities at M12').bold = True
        row.cells[2].paragraphs[0].add_run('Planned functionalities at M18').bold = True
        row.cells[3].paragraphs[0].add_run('Planned functionalities at M24').bold = True
        
    else:
        print '[Warning]: only stype == (0 or 1) supported'


    for i in range(1,num_rows+1):
        #
        # 1 : M18
        # 2 : M12
        # 3 : M24
    
        table.rows[i].cells[0].text = summary_data[i-1][0] # Component
        table.rows[i].cells[1].text = summary_data[i-1][2]
        table.rows[i].cells[2].text = summary_data[i-1][1]
        table.rows[i].cells[3].text = summary_data[i-1][3]
   
    if stype == 0:
        document.save('summary_components_build_prod_comp_task.docx')    
    elif stype == 1:
        document.save('summary_components_planned_functionality.docx')            

    
    


