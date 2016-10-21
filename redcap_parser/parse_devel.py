import redcap_parser as rp
import data2docx as d2d
import data2excel as d2e
from openpyxl import Workbook


from docx import Document
from docx.shared import Pt
import utilities as util
import parse as pp
if __name__ == '__main__':
    
    
    comps       = pp.parse('/home/guillaume/Documents/redcap/data.csv')
    all_tasks   = util.get_all_tasks(comps)

    stype=0

    summary_data = util.summary(comps,stype)
    
#%%    
    
    d2d.summary_comp2docx(summary_data,stype)
    
    #%%

    deps = util.get_list_depenencies(summary_data)    
    
    
    wb = d2e.comp2excel(summary_data,deps)
    
    
#%% function    
    
        
def get_dep_indicies(dep,deps):
    """
        For a component name find the row indicies of the dependencies.
        
        Parameters
        ----------
        
        dep: list of dependecies: ['Image & Genetic Viewer',..]
        
    """
    
    indicies = []
 
    for d in dep:
        for i in range(0,len(deps)):
            if d == deps[i]:
                indicies.extend([i])
            
    return indicies
        
        
    
#%% DEVEL

    
    wb  = Workbook()
    ws1 = wb.create_sheet("summary",0) #

    # Add dependency names in columne
    r = 2    
    for dep in deps:
        ws1.cell(row=r, column=1, value=dep)    
        r = r + 1


    # Add dependency names in columne
    col = 2    
    for summ in summary_data:
        ws1.cell(row=1, column=col, value=summ['name']) 
        
        print summ
        
        idx = get_dep_indicies(summ['dep'],deps)
        
        print idx
        
        if idx:
            for i in idx:
                cs = ws1.cell(row=i+2, column=col, value='1')    
                cs.style.alignment.horizontal = 'center'
        
        col = col + 1
            
    
#%%    
    
    wb.save('summary.xlsx')
    
    
    #%%
    d2d.summary_comp2docx(summary_data,stype)


    len(summary_data)

    summary_data[2]






#%% Check if have component

data_tnames         = ['T8.2.1','T8.4.1']
data_factor_tnames  = ['T8.1.1','T8.5.2',''] 

task_names = ['T8.3.11',]

has_v = util.has_tasks(summary_data,task_names,3)


has_v

#%% PUT TASKS FROM ALL COMPONENTS TOGTHER AND SAVE


print 'total number of tasks: ', len(all_tasks)
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

tree = d2d.sorted_task2docx(document,all_tasks)

document.save('redcap_sorted.docx')    
    
    
#%%

test = [['18', 'First dataset annotated according to the ontology.'], ['12', 'Prototype of the ontology for describing data on patients with neurological diseases developed'], ['24', 'Ontology and several datasets annotated according to the ontology.']]


#%% SAVE ONE COMPONENT (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

row_id              = 2
col_names, data_val = rp.extrat_row(ws,row_id)   
comp                = util.row2comp(col_names,data_val)

print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'

comps.append(comp)

d2d.comp2docx(document,comp,bSorted=True)   
    
document.save('test4.docx')    

    
#%% SAVE ALL COMPONENTS (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

i = 0
for comp in comps:
    print '=====> ', i, '<====='
    d2d.comp2docx(document,comp,bSorted=False)   
    i = i + 1

document.save('test3.docx')    











import redcap_parser as rp
import data2docx as d2d
from docx import Document
from docx.shared import Pt
import utilities as util

if __name__ == '__main__':
    
    filename = "/home/guillaume/Documents/redhat/data3.csv"
    data     = rp.import_csv(filename)
        
    
    row_id = 0
    col_names, data_val = rp.csv2data(data,row_id)
    comp                = util.row2comp(col_names,data_val)


#%%

# Extracting components

num_rows            = len(data)-1
comps               = []


for row_id in range(0,num_rows):
    
    col_names, data_val = rp.csv2data(data,row_id)
    comp                = util.row2comp(col_names,data_val)
    
    print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'
    comps.append(comp)
       
    
#%% Print

print ' '
for comp in comps:
    print comp    
    print ' '
    
#%% import csv
