#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  3 17:24:32 2016

@author: guillaume
"""

from docx import Document
import redcap_parser as rp


def one_line_2_docx(p,name,value=''):
    """
        Prints to word document format:
            
            <b>name</b>value\n
    
    """
    # print component    
    p.add_run(name).bold = True
    if value != '':
        p.add_run(value)
        p.add_run('\n')
    

def dependencies2docx(wdoc,data_hosp,data_ref,soft):
    """
        
    Takes Data and software information and puts it into a table
    
    Prameters
    ---------
     
    wdoc: python-docx word document object
    data_hosp: DATA > Hospital Data > [] list
    data_ref : DATA > Reference Data > [] list

    """
    
    data_hosp = ', '.join(data_hosp)
    data_ref  = ', '.join(data_ref)
    
    
    if len(data_hosp) == 0:
        bDataHosp = False
    else:
        bDataHosp = True
        
    if len(data_ref) == 0:
        bDataRef = False
    else:
        bDataRef = True
        
        
   
    # get number of sofware componets
    num_soft = len(soft.keys())
    
    tcase = 'full'
    
    if bDataHosp & bDataRef  :
        tab_num_rows = 2 + num_soft
        tcase        = 'full'    
    elif (bDataHosp == True) & (bDataRef == False):
        tab_num_rows = 1 + num_soft
        tcase        = 'only_hosp'         
    elif (bDataHosp == False) & (bDataRef == True):
        tab_num_rows = 1 + num_soft
        tcase        = 'only_ref'         
    else:
        tab_num_rows = num_soft   
        tcase        = 'only_soft'         
    
        
    if tab_num_rows == 0:
        return
            
    table       = wdoc.add_table(rows=tab_num_rows, cols=3)

    
    

    if tcase == 'full':
        
        # Hospital Data
        row = table.rows[0]
        row.cells[0].text = 'DATA'
        row.cells[1].text = 'Hospital'
        row.cells[2].text = data_hosp

        # Reference Data
        row = table.rows[1]
        row.cells[1].text = 'Reference'
        row.cells[2].text = data_ref

        r_idx = 2

    elif tcase == 'only_hosp':
        
        # Hospital Data
        row = table.rows[0]
        row.cells[0].text = 'DATA'
        row.cells[1].text = 'Hospital'
        row.cells[2].text = data_hosp

        r_idx = 1
        
    elif tcase == 'only_ref':

        # Reference Data
        row = table.rows[0]
        row.cells[1].text = 'Reference'
        row.cells[2].text = data_ref

        r_idx = 1
    
    else:
    
        r_idx = 0
        
    
    print tcase, r_idx, tab_num_rows
        

    # Software Data
    row = table.rows[r_idx]
    row.cells[0].text = 'SOFTWARE'
    

    idx = r_idx    
    for key in soft:
        row = table.rows[idx]
        soft_desc = ', '.join(soft[key])
        row.cells[1].text = key
        row.cells[2].text = soft_desc
        idx = idx + 1
    



def releases2docx(wdoc,release):
    """
        Takes release information  (Planned functional etc..)
    
    
    """
    
    num_releases = len(release.keys())
       
    table = wdoc.add_table(rows=num_releases, cols=2)
    
    idx = 0    
    for key in release:
        row = table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = release[key]

        idx = idx + 1

def usecase2docx(wdoc,use_case):
    
    
    table = wdoc.add_table(rows=2, cols=2)

    # Heading
    row = table.rows[0]
    row.cells[0].text = 'User'
    row.cells[1].text = 'Use cases'

    # User & Use cases    
    row = table.rows[1]
    row.cells[0].text = use_case['users']
    row.cells[1].text = use_case['desc']
    

    
def component2docx(document,comp):
    
    
    
    print 'comp.titles: ', comp.titles
    
    for i in range(0,len(comp.titles)):
        document.add_heading(comp.titles[i], level=i+1)
        
    
    # Titles
    #document.add_heading(comp.titles[0], level=1)
    #document.add_heading(comp.titles[1], level=2)
    #document.add_heading(comp.titles[2], level=3)
    
    
    p = document.add_paragraph()
    
    one_line_2_docx(p,'Component(' + comp.subject_ID  + '): ',            comp.component)
    one_line_2_docx(p,'Contributing task(s): ', comp.contribution)
    one_line_2_docx(p,'Description: ',          comp.short_desc)
    
    
    # Prior Dependencies table:
    p = document.add_paragraph()
    p.style = 'Normal'
    
    one_line_2_docx(p,'Dependencies:')
    
    #  add table
    dependencies2docx(document,comp.data_hosp,comp.data_ref,comp.soft)

    # Prior Release table:
    p = document.add_paragraph()
    p.add_run('\n')
    one_line_2_docx(p,'Releases:')
    releases2docx(document,comp.release)
    
    
    # Prior User table:
    p = document.add_paragraph()
    p.add_run('\n')
    one_line_2_docx(p,'User and Use cases:')   
    usecase2docx(document,comp.use_case)
    



def data2word(data_vals,col_names,document):
    """ 
        data_vals   : list 
        
        col_names   : list
    
        document    : word document 
    
    """
    subject_ID = data_vals[0]
    document.add_heading('Subject: ' + subject_ID, level=3)
    p = document.add_paragraph('')
    run = p.add_run()
    run.add_break()    

    for i in range(1,len(col_names)):
        col_name = col_names[i]
        data_val = data_vals[i]
        

        if data_val != 'None':
            run = p.add_run()
            run.font.bold = True
            run.add_text(col_name + '\n')
            run = p.add_run()
            run.font.bold = False
            run.add_text(data_val + '\n\n')

            
            run.add_break()    
   

def save2word_document(ws):
    document = Document()

    document.add_heading('redcap_data', 0)
    document.add_paragraph('automatically generated from python')

    num_rows = ws.max_row

    for i in range(2,num_rows):
        col_names,data_val = rp.extrat_row(ws,i)
        print data_val
        data2word(data_val,col_names,document)

    document.save('redcap_python.docx')    
    return data_val, col_names
            