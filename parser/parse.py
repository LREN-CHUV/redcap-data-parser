import redcap_parser as rp
import data2docx as d2d
from docx import Document
from docx.shared import Pt
import utilities as util


if __name__ == '__main__':
    
    filename = "/home/guillaume/Documents/redhat/data.xlsx"
    ws       = rp.import_excel(filename)
       
    


# Extracting components

num_rows            = ws.max_row
comps               = []


for row_id in range(2,num_rows):
    
    col_names, data_val = rp.extrat_row(ws,row_id)   
    comp                = util.row2comp(col_names,data_val)
    
    print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'
    comps.append(comp)
    
    

#%% PUT TASKS FROM ALL COMPONENTS TOGTHER AND SAVE


all_tasks = []

for comp in comps:
    print 'Subject: ', comp.subject_ID, ' num_tasks: ', len(comp.tasks)
    all_tasks.extend(comp.tasks)


print 'total number of tasks: ', len(all_tasks)


document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

d2d.sorted_task2docx(document,all_tasks)

document.save('redcap_sorted.docx')    
    
   

#%% SAVE ONE COMPONENT (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

row_id              = 6
col_names, data_val = rp.extrat_row(ws,row_id)   
comp                = util.row2comp(col_names,data_val)

print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'

comps.append(comp)

d2d.comp2docx(document,comp,bSorted=True)   
    
document.save('test4.docx')    

    
#%% SAVE ALL COMPONENTS (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

i = 0
for comp in comps:
    print '=====> ', i, '<====='
    d2d.comp2docx(document,comp,bSorted=True)   
    i = i + 1

document.save('test3.docx')    







